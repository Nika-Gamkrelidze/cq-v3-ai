"""`extract_text` against the files tenants actually upload.

Two production incidents drove these guards, both found by importing a replica of a real
customer document (a Georgian bank's installment-terms DOCX, then a call-QA scorecard XLSX):

  * DOCX tables were silently dropped — `document.paragraphs` does not contain them, so the
    load-bearing facts (amounts, terms, rates) never reached the knowledge base while the
    import still reported ready.
  * Unknown binary uploads (an .xlsx before support existed) fell through to a utf-8 decode
    and died much later inside Postgres as `invalid byte sequence for encoding "UTF8"` —
    an error message the uploader cannot act on.

Everything here is in-memory: openpyxl/python-docx build the fixture bytes, extract_text
parses them back. No database, no event loop.
"""
import io

import pytest

from app.services.kb_ingest import extract_text


def _xlsx_bytes() -> bytes:
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "შეფასება"
    ws.append(["შემოწმების თარიღი", ""])                       # label row, no value
    ws.append(["A1", "თანამშრომელმა უპასუხა ზარს", "1/-3", 1])  # criterion row
    ws.append([None, None, None])                               # fully empty row
    ws.append(["B7", "ლოდინი ხაზზე", "#REF!", "#REF!"])         # cached formula errors
    ws2 = wb.create_sheet("წონები")
    ws2.append(["სექცია", "მაქს. ქულა"])
    ws2.append(["A", 10])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _docx_bytes() -> bytes:
    import docx
    d = docx.Document()
    d.add_paragraph("სესხის პირობები")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "სესხის თანხა"
    t.cell(0, 1).text = "100-დან 7,000 ლარამდე"
    t.cell(1, 0).text = "ვადა"
    t.cell(1, 1).text = "1-დან 36 თვემდე"
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_xlsx_rows_survive_and_formula_errors_do_not():
    text = extract_text("scorecard.xlsx", "application/vnd.openxmlformats-officedocument"
                        ".spreadsheetml.sheet", _xlsx_bytes())
    assert "თანამშრომელმა უპასუხა ზარს" in text
    assert "1/-3" in text                    # weights are facts, they stay
    assert "#REF!" not in text               # cached formula errors are noise, they go
    assert "[წონები]" in text                # multi-sheet workbooks name their sheets
    assert "მაქს. ქულა" in text


def test_docx_tables_are_extracted():
    text = extract_text("terms.docx", "application/vnd.openxmlformats-officedocument"
                        ".wordprocessingml.document", _docx_bytes())
    assert "სესხის პირობები" in text                      # paragraph
    assert "სესხის თანხა | 100-დან 7,000 ლარამდე" in text  # table row, flattened
    assert "ვადა | 1-დან 36 თვემდე" in text


@pytest.mark.parametrize("name,must_mention", [
    ("old.xls", ".xlsx"),
    ("old.doc", ".docx"),
])
def test_legacy_office_formats_get_actionable_errors(name, must_mention):
    with pytest.raises(ValueError) as e:
        extract_text(name, "application/octet-stream", b"\xd0\xcf\x11\xe0" + b"\x00" * 64)
    assert must_mention in str(e.value)


def test_unknown_binary_is_rejected_not_decoded():
    blob = bytes(range(256)) * 32           # dense non-UTF-8 binary
    with pytest.raises(ValueError) as e:
        extract_text("photo.png", "image/png", blob)
    assert "Supported formats" in str(e.value)


def test_plain_text_still_passes_through_with_nuls_stripped():
    text = extract_text("notes.txt", "text/plain", "გამარჯობა\x00 მსოფლიო".encode())
    assert text == "გამარჯობა მსოფლიო"
